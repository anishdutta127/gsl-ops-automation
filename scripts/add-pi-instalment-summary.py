"""Inject the Phase 5 instalment summary table into pi-template.docx.

Idempotent: if a paragraph already contains "Instalment Summary" we
exit without touching anything. Otherwise we append (in document order):

  1. A "Instalment Summary" heading paragraph.
  2. A 4-column table with one header row + one loop row using the
     docxtemplater `{#INSTALMENT_SUMMARY}...{/INSTALMENT_SUMMARY}`
     row-loop pattern (same shape the existing LINE_ITEMS table uses).
  3. Two summary paragraphs that consume CURRENT_STUDENT_COUNT,
     CONTRACT_TOTAL_AT_CURRENT_COUNT, and TOTAL_RECEIVED_TO_DATE.

The insertion lands AFTER the line-items table and BEFORE the
{PAYMENT_TERMS} paragraph so the summary reads in narrative order on
the PI: line items first, then schedule context, then payment terms.

The placeholders match the bag produced by
src/lib/pi/generatePi.ts:buildPlaceholderBag.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from copy import deepcopy

TEMPLATE = Path("public/ops-templates/pi-template.docx")


def already_patched(doc) -> bool:
    for p in doc.paragraphs:
        if "Instalment Summary" in p.text:
            return True
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "INSTALMENT_SUMMARY" in cell.text:
                    return True
    return False


def main() -> int:
    if not TEMPLATE.exists():
        print(f"Template not found: {TEMPLATE}", file=sys.stderr)
        return 1

    doc = Document(str(TEMPLATE))
    if already_patched(doc):
        print("Template already patched; no changes made.")
        return 0

    # Locate the line-items table (the one with {#LINE_ITEMS}) and
    # the PAYMENT_TERMS paragraph (the anchor we insert before).
    line_items_table = None
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                if "{#LINE_ITEMS}" in cell.text:
                    line_items_table = t
                    break
            if line_items_table is not None:
                break
        if line_items_table is not None:
            break
    if line_items_table is None:
        print("Could not find LINE_ITEMS table to anchor after.", file=sys.stderr)
        return 2

    payment_terms_para = None
    for p in doc.paragraphs:
        if "{PAYMENT_TERMS}" in p.text:
            payment_terms_para = p
            break
    if payment_terms_para is None:
        print("Could not find {PAYMENT_TERMS} paragraph anchor.", file=sys.stderr)
        return 3

    # Build the new heading paragraph.
    heading_p = doc.paragraphs[0]._element.makeelement(qn("w:p"), {})
    pPr = heading_p.makeelement(qn("w:pPr"), {})
    pStyle = pPr.makeelement(qn("w:pStyle"), {qn("w:val"): "Heading2"})
    pPr.append(pStyle)
    heading_p.append(pPr)
    r = heading_p.makeelement(qn("w:r"), {})
    rPr = r.makeelement(qn("w:rPr"), {})
    b = rPr.makeelement(qn("w:b"), {})
    rPr.append(b)
    r.append(rPr)
    t_el = r.makeelement(qn("w:t"), {})
    t_el.text = "Instalment Summary"
    r.append(t_el)
    heading_p.append(r)

    # Build the new 4-column table.
    new_table_xml = """
    <w:tbl xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:tblPr>
        <w:tblStyle w:val="TableGrid"/>
        <w:tblW w:w="0" w:type="auto"/>
        <w:tblBorders>
          <w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
          <w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
        </w:tblBorders>
      </w:tblPr>
      <w:tblGrid>
        <w:gridCol w:w="700"/>
        <w:gridCol w:w="2200"/>
        <w:gridCol w:w="3200"/>
        <w:gridCol w:w="2700"/>
      </w:tblGrid>
      <w:tr>
        <w:tc>
          <w:tcPr><w:tcW w:w="700" w:type="dxa"/></w:tcPr>
          <w:p><w:r><w:rPr><w:b/></w:rPr><w:t>#</w:t></w:r></w:p>
        </w:tc>
        <w:tc>
          <w:tcPr><w:tcW w:w="2200" w:type="dxa"/></w:tcPr>
          <w:p><w:r><w:rPr><w:b/></w:rPr><w:t>Due date</w:t></w:r></w:p>
        </w:tc>
        <w:tc>
          <w:tcPr><w:tcW w:w="3200" w:type="dxa"/></w:tcPr>
          <w:p><w:r><w:rPr><w:b/></w:rPr><w:t>Status</w:t></w:r></w:p>
        </w:tc>
        <w:tc>
          <w:tcPr><w:tcW w:w="2700" w:type="dxa"/></w:tcPr>
          <w:p><w:r><w:rPr><w:b/></w:rPr><w:t>Amount</w:t></w:r></w:p>
        </w:tc>
      </w:tr>
      <w:tr>
        <w:tc>
          <w:tcPr><w:tcW w:w="700" w:type="dxa"/></w:tcPr>
          <w:p><w:r><w:t>{#INSTALMENT_SUMMARY}{seq}</w:t></w:r></w:p>
        </w:tc>
        <w:tc>
          <w:tcPr><w:tcW w:w="2200" w:type="dxa"/></w:tcPr>
          <w:p><w:r><w:t>{dueDate}</w:t></w:r></w:p>
        </w:tc>
        <w:tc>
          <w:tcPr><w:tcW w:w="3200" w:type="dxa"/></w:tcPr>
          <w:p><w:r><w:t>{status}</w:t></w:r></w:p>
          <w:p><w:r><w:rPr><w:i/><w:sz w:val="18"/></w:rPr><w:t>{breakdown}</w:t></w:r></w:p>
        </w:tc>
        <w:tc>
          <w:tcPr><w:tcW w:w="2700" w:type="dxa"/></w:tcPr>
          <w:p><w:r><w:t>{amount}{/INSTALMENT_SUMMARY}</w:t></w:r></w:p>
        </w:tc>
      </w:tr>
    </w:tbl>
    """.strip()

    from lxml import etree
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    new_table_el = etree.fromstring(new_table_xml)

    # Build the two summary paragraphs.
    def make_summary_paragraph(text: str):
        p_xml = f"""
        <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
          <w:pPr><w:spacing w:before="80" w:after="40"/></w:pPr>
          <w:r><w:t xml:space="preserve">{text}</w:t></w:r>
        </w:p>
        """.strip()
        return etree.fromstring(p_xml)

    summary1 = make_summary_paragraph(
        "Contract total at {CURRENT_STUDENT_COUNT} students: {CONTRACT_TOTAL_AT_CURRENT_COUNT}"
    )
    summary2 = make_summary_paragraph(
        "Total received to date: {TOTAL_RECEIVED_TO_DATE}"
    )

    # Insertion strategy: insert the heading, table, and summary
    # paragraphs immediately BEFORE the PAYMENT_TERMS paragraph
    # element. python-docx exposes _element.addprevious which mirrors
    # what the underlying lxml tree supports.
    anchor = payment_terms_para._element
    # Walk up to the body if PAYMENT_TERMS is wrapped (it's not, but
    # defensive).
    body = anchor.getparent()
    insert_index = list(body).index(anchor)
    body.insert(insert_index, heading_p)
    body.insert(insert_index + 1, new_table_el)
    body.insert(insert_index + 2, summary1)
    body.insert(insert_index + 3, summary2)

    doc.save(str(TEMPLATE))
    print(f"Patched {TEMPLATE}: added heading + INSTALMENT_SUMMARY table + 2 summary paragraphs.")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
